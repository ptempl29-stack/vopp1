import uuid
import io
import re

from fastapi import APIRouter, HTTPException, Depends, UploadFile, File, Form
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel, EmailStr
from typing import Optional

from core.db import db, now_iso, logger
from core.config import FORMS_ROLES, FORM_STATUSES, EXT_CONTENT_TYPES, APP_NAME, PUBLIC_BASE_URL
from core.security import require_roles
from core.audit import log_audit
from core.storage import put_object, get_object
from core.email_utils import send_email, email_configured
from core.sms_utils import send_sms, sms_configured
from core.pdf_utils import new_pdf, pdf_bytes, FONT
from models.schemas import FormInput, FormSubmission, IdList
from data.seed import FORM_TEMPLATES

router = APIRouter()

_TEMPLATE_META = {
    "docx": ("application/vnd.openxmlformats-officedocument.wordprocessingml.document", "blank_form.docx"),
    "xlsx": ("application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", "blank_spreadsheet.xlsx"),
    "pdf": ("application/pdf", "blank_form.pdf"),
}


def _build_docx(clinic: str) -> bytes:
    from docx import Document
    doc = Document()
    doc.add_heading(clinic, level=0)
    doc.add_heading("Patient Form", level=1)
    for label in ["Patient Name:", "Date of Birth:", "Date:", "Notes:"]:
        doc.add_paragraph(label)
    for _ in range(12):
        doc.add_paragraph("")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_xlsx(clinic: str) -> bytes:
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Form"
    ws["A1"] = clinic
    ws["A2"] = "Patient Form"
    ws.append([])
    ws.append(["Field", "Value"])
    for label in ["Patient Name", "Date of Birth", "Date", "Notes"]:
        ws.append([label, ""])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def _build_pdf(clinic: str) -> bytes:
    from core.pdf_utils import new_pdf, pdf_bytes, FONT
    pdf = new_pdf()
    pdf.set_font(FONT, "B", 18)
    pdf.cell(0, 12, clinic, ln=True)
    pdf.set_font(FONT, "B", 14)
    pdf.cell(0, 10, "Patient Form", ln=True)
    pdf.ln(4)
    pdf.set_font(FONT, "", 12)
    for label in ["Patient Name: ______________________________", "Date of Birth: _____________________________",
                  "Date: ______________________________________", "Notes:"]:
        pdf.cell(0, 10, label, ln=True)
    for _ in range(10):
        pdf.cell(0, 10, "_" * 70, ln=True)
    return pdf_bytes(pdf)


@router.get("/forms/blank-template/{kind}")
async def blank_template(kind: str, user: dict = Depends(require_roles(*FORMS_ROLES))):
    if kind not in _TEMPLATE_META:
        raise HTTPException(status_code=400, detail="Unsupported template type")
    s = await db.settings.find_one({"key": "clinic"}, {"_id": 0})
    clinic = (s or {}).get("clinic_name", "Veterans of Puerto Plata")
    builders = {"docx": _build_docx, "xlsx": _build_xlsx, "pdf": _build_pdf}
    data = builders[kind](clinic)
    ctype, fname = _TEMPLATE_META[kind]
    return StreamingResponse(io.BytesIO(data), media_type=ctype,
        headers={"Content-Disposition": f'attachment; filename="{fname}"'})


@router.get("/forms")
async def list_forms(user: dict = Depends(require_roles("doctor", "nurse", "receptionist"))):
    forms = await db.forms.find({}, {"_id": 0}).sort("created_at", -1).to_list(500)
    patients = {p["id"]: f"{p['first_name']} {p['last_name']}" async for p in db.patients.find({}, {"_id": 0})}
    for f in forms:
        f["patient_name"] = patients.get(f.get("patient_id"), "-")
    await log_audit("view", "form", actor=user, detail=f"list ({len(forms)})")
    return forms


@router.post("/forms")
async def create_form(data: FormInput, user: dict = Depends(require_roles(*FORMS_ROLES))):
    doc = {"id": str(uuid.uuid4()), "public_token": uuid.uuid4().hex,
           "patient_id": data.patient_id, "title": data.title, "form_type": data.form_type,
           "fields": data.fields, "external_url": data.external_url, "status": data.status,
           "template": FORM_TEMPLATES.get(data.form_type, []), "responses": None,
           "attachment": None, "email_sent": False,
           "created_at": now_iso(), "created_by": user["name"]}
    recipient = data.recipient_email
    if not recipient and data.patient_id:
        p = await db.patients.find_one({"id": data.patient_id}, {"_id": 0})
        if p and p.get("email"):
            recipient = p["email"]
    if recipient:
        link = f"{PUBLIC_BASE_URL.rstrip('/')}/form/{doc['public_token']}"
        sent = send_email(recipient,
            f"Please complete your form: {data.title}",
            f"Hello,\n\nPlease complete the following form for Veterans of Puerto Plata:\n{data.title}\n\n{link}\n\nThank you.")
        doc["email_sent"] = sent
        doc["recipient_email"] = recipient
    await db.forms.insert_one(doc)
    doc.pop("_id", None)
    await log_audit("create", "form", actor=user, resource_id=doc["id"],
                    detail=f"{data.form_type}: {data.title}")
    return doc


@router.post("/forms/upload")
async def upload_form(file: UploadFile = File(...), title: str = Form(...),
                      form_type: str = Form("Uploaded"), patient_id: str = Form(""),
                      user: dict = Depends(require_roles(*FORMS_ROLES))):
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else "bin"
    if ext not in EXT_CONTENT_TYPES:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    safe_ct = EXT_CONTENT_TYPES[ext]
    path = f"{APP_NAME}/forms/{user['id']}/{uuid.uuid4()}.{ext}"
    payload = await file.read()
    if len(payload) > 15 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 15MB)")
    try:
        result = put_object(path, payload, safe_ct)
    except Exception as e:
        logger.error(f"upload failed: {e}")
        raise HTTPException(status_code=502, detail="File storage failed")
    doc = {"id": str(uuid.uuid4()), "public_token": uuid.uuid4().hex,
           "title": title, "form_type": form_type, "patient_id": patient_id or None,
           "template": [], "responses": None, "status": "received",
           "attachment": {"storage_path": result["path"], "filename": file.filename,
                          "content_type": safe_ct, "size": result.get("size")},
           "created_at": now_iso(), "created_by": user["name"]}
    await db.forms.insert_one(doc)
    doc.pop("_id", None)
    await log_audit("create", "form", actor=user, resource_id=doc["id"],
                    detail=f"upload: {file.filename}")
    return doc


@router.get("/forms/{fid}/download")
async def download_form(fid: str, user: dict = Depends(require_roles(*FORMS_ROLES))):
    f = await db.forms.find_one({"id": fid})
    if not f or not f.get("attachment"):
        raise HTTPException(status_code=404, detail="Attachment not found")
    data, ctype = get_object(f["attachment"]["storage_path"])
    fname = f["attachment"]["filename"].replace('"', "").replace("\n", "").replace("\r", "")
    await log_audit("view", "form", actor=user, resource_id=fid, detail=f"download: {fname}")
    return Response(content=data, media_type="application/octet-stream",
        headers={"Content-Disposition": f'attachment; filename="{fname}"',
                 "X-Content-Type-Options": "nosniff"})


@router.post("/forms/{fid}/send-sms")
async def send_form_sms(fid: str, user: dict = Depends(require_roles(*FORMS_ROLES))):
    f = await db.forms.find_one({"id": fid}, {"_id": 0})
    if not f:
        raise HTTPException(status_code=404, detail="Form not found")
    if not f.get("public_token"):
        raise HTTPException(status_code=400, detail="This form has no shareable link")
    if not sms_configured():
        return {"sent": False, "configured": False}
    phone = None
    if f.get("patient_id"):
        p = await db.patients.find_one({"id": f["patient_id"]}, {"_id": 0, "phone": 1})
        phone = (p or {}).get("phone")
    if not phone:
        raise HTTPException(status_code=400, detail="Patient has no phone number on file")
    s = await db.settings.find_one({"key": "clinic"}, {"_id": 0})
    clinic = (s or {}).get("clinic_name", "Veterans of Puerto Plata")
    link = f"{PUBLIC_BASE_URL.rstrip('/')}/form/{f['public_token']}"
    ok = send_sms(phone, f'{clinic}: Please complete your form "{f["title"]}": {link}')
    await log_audit("update", "form", actor=user, resource_id=fid, detail=f"sms sent={ok}")
    return {"sent": ok, "configured": True}


@router.put("/forms/{fid}/status")
async def update_form_status(fid: str, status: str, user: dict = Depends(require_roles(*FORMS_ROLES))):
    if status not in FORM_STATUSES:
        raise HTTPException(status_code=400, detail="Invalid status")
    res = await db.forms.update_one({"id": fid}, {"$set": {"status": status}})
    if res.matched_count == 0:
        raise HTTPException(status_code=404, detail="Form not found")
    await log_audit("update", "form", actor=user, resource_id=fid, detail=f"status={status}")
    return await db.forms.find_one({"id": fid}, {"_id": 0})


class FormUpdate(BaseModel):
    title: Optional[str] = None
    form_type: Optional[str] = None
    status: Optional[str] = None
    patient_id: Optional[str] = None
    external_url: Optional[str] = None
    recipient_email: Optional[str] = None
    responses: Optional[dict] = None


class EmailSend(BaseModel):
    recipients: Optional[list] = None
    recipient_email: Optional[EmailStr] = None


class ToFolder(BaseModel):
    patient_id: str
    subfolder_id: Optional[str] = None


async def _patient_name(pid):
    if not pid:
        return None
    p = await db.patients.find_one({"id": pid}, {"_id": 0, "first_name": 1, "last_name": 1})
    return f"{p['first_name']} {p['last_name']}" if p else None


def _form_pdf(form: dict, clinic: str) -> bytes:
    pdf = new_pdf()
    pdf.set_font(FONT, "B", 18)
    pdf.cell(0, 12, (clinic or "")[:60], ln=True)
    pdf.set_font(FONT, "B", 13)
    pdf.cell(0, 9, (form.get("title") or "Patient Form")[:70], ln=True)
    pdf.set_font(FONT, "", 10)
    meta = " · ".join([x for x in [form.get("form_type"), form.get("patient_name")] if x])
    if meta:
        pdf.cell(0, 7, meta, ln=True)
    pdf.ln(3)
    responses = form.get("responses") or {}
    template = form.get("template") or []

    def _mc(text, bold=False, size=11):
        pdf.set_font(FONT, "B" if bold else "", size)
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(0, 6, text)

    if template and responses:
        for fld in template:
            val = responses.get(fld.get("name"))
            if val in (None, ""):
                continue
            _mc(str(fld.get("en") or fld.get("name")), bold=True, size=10)
            shown = "[signature on file]" if str(val).startswith("data:image") else str(val)
            _mc(shown)
            pdf.ln(1)
    elif responses:
        for k, v in responses.items():
            _mc(str(k), bold=True, size=10)
            shown = "[signature on file]" if str(v).startswith("data:image") else str(v)
            _mc(shown)
            pdf.ln(1)
    else:
        _mc("No responses recorded for this form yet.")
    return pdf_bytes(pdf)


@router.put("/forms/{fid}")
async def update_form(fid: str, data: FormUpdate, user: dict = Depends(require_roles(*FORMS_ROLES))):
    existing = await db.forms.find_one({"id": fid}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Form not found")
    if user["role"] != "admin" and existing.get("created_by") not in (user["name"], None):
        raise HTTPException(status_code=403, detail="You can only edit forms you created")
    updates = {}
    for k in ("title", "form_type", "patient_id", "external_url", "recipient_email", "responses"):
        v = getattr(data, k)
        if v is not None:
            updates[k] = v
    if data.status is not None:
        if data.status not in FORM_STATUSES:
            raise HTTPException(status_code=400, detail="Invalid status")
        updates["status"] = data.status
    if updates:
        updates["updated_at"] = now_iso()
        updates["updated_by"] = user["name"]
        await db.forms.update_one({"id": fid}, {"$set": updates})
    await log_audit("update", "form", actor=user, resource_id=fid, detail="edit form")
    return await db.forms.find_one({"id": fid}, {"_id": 0})


@router.post("/forms/{fid}/send-email")
async def send_form_email(fid: str, data: EmailSend, user: dict = Depends(require_roles(*FORMS_ROLES))):
    f = await db.forms.find_one({"id": fid})
    if not f:
        raise HTTPException(status_code=404, detail="Form not found")
    raw = list(data.recipients or [])
    if data.recipient_email:
        raw.append(str(data.recipient_email))
    seen, recipients = set(), []
    for r in raw:
        r = (r or "").strip()
        if r and re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", r) and r.lower() not in seen:
            seen.add(r.lower())
            recipients.append(r)
    if not recipients:
        raise HTTPException(status_code=400, detail="At least one valid recipient email is required")
    if len(recipients) > 20:
        raise HTTPException(status_code=400, detail="Too many recipients (max 20 per send)")
    if not email_configured():
        return {"sent": 0, "total": len(recipients), "configured": False, "failed": recipients}
    s = await db.settings.find_one({"key": "clinic"}, {"_id": 0})
    clinic = (s or {}).get("clinic_name", "Veterans of Puerto Plata")
    body_lines = ["Hello,", "", f"{clinic} has sent you a form: {f.get('title')}"]
    if f.get("public_token"):
        link = f"{PUBLIC_BASE_URL.rstrip('/')}/form/{f['public_token']}"
        body_lines += ["", "Please complete it securely here:", link]
    if f.get("external_url"):
        body_lines += ["", "Reference link:", f["external_url"]]
    body_lines += ["", "Thank you."]
    attachments = None
    if f.get("attachment"):
        try:
            content, _ = get_object(f["attachment"]["storage_path"])
            attachments = [{"filename": f["attachment"]["filename"], "content": content,
                            "content_type": f["attachment"].get("content_type", "application/octet-stream")}]
        except Exception as e:
            logger.error(f"attach fetch failed: {e}")
    subject = f"{clinic}: {f.get('title')}"
    body = "\n".join(body_lines)
    sent, failed = 0, []
    for r in recipients:
        if send_email(r, subject, body, attachments):
            sent += 1
        else:
            failed.append(r)
    await db.forms.update_one({"id": fid}, {"$set": {"email_sent": sent > 0, "recipient_email": recipients[0]}})
    await log_audit("update", "form", actor=user, resource_id=fid,
                    detail=f"email sent {sent}/{len(recipients)} to {', '.join(recipients)}")
    return {"sent": sent, "total": len(recipients), "configured": True, "failed": failed}


@router.post("/forms/{fid}/to-folder")
async def move_form_to_folder(fid: str, data: ToFolder, user: dict = Depends(require_roles(*FORMS_ROLES))):
    f = await db.forms.find_one({"id": fid}, {"_id": 0})
    if not f:
        raise HTTPException(status_code=404, detail="Form not found")
    pname = await _patient_name(data.patient_id)
    if not pname:
        raise HTTPException(status_code=404, detail="Patient not found")
    if data.subfolder_id:
        sf = await db.folder_subfolders.find_one({"id": data.subfolder_id, "patient_id": data.patient_id})
        if not sf:
            raise HTTPException(status_code=400, detail="Target folder not found for that patient")
    base = {"id": str(uuid.uuid4()), "patient_id": data.patient_id, "subfolder_id": data.subfolder_id or None,
            "form_id": f["id"], "description": f.get("form_type") or "",
            "created_at": now_iso(), "created_by": user["name"], "created_by_id": user["id"]}
    if f.get("attachment"):
        att = f["attachment"]
        item = {**base, "source": "form", "storage_path": att["storage_path"], "filename": att["filename"],
                "label": (f.get("title") or att["filename"])[:120],
                "content_type": att.get("content_type", "application/octet-stream"), "size": att.get("size")}
    else:
        s = await db.settings.find_one({"key": "clinic"}, {"_id": 0})
        clinic = (s or {}).get("clinic_name", "Veterans of Puerto Plata")
        form_with_name = {**f, "patient_name": pname}
        try:
            pdf = _form_pdf(form_with_name, clinic)
        except Exception as e:
            logger.error(f"form pdf failed: {e}")
            raise HTTPException(status_code=500, detail="Could not render form PDF")
        fname = f"{(f.get('title') or 'form').replace('/', '-')[:50]}.pdf"
        path = f"{APP_NAME}/folders/{data.patient_id}/{uuid.uuid4()}.pdf"
        try:
            result = put_object(path, pdf, "application/pdf")
        except Exception as e:
            logger.error(f"form pdf storage failed: {e}")
            raise HTTPException(status_code=502, detail="File storage failed")
        item = {**base, "source": "upload", "storage_path": result["path"], "filename": fname,
                "label": (f.get("title") or "Form")[:120], "content_type": "application/pdf",
                "size": result.get("size")}
    await db.folder_items.insert_one(item)
    item.pop("_id", None)
    await log_audit("update", "form", actor=user, resource_id=fid, detail=f"to folder {pname}")
    return {"ok": True, "patient_name": pname}


# ---------------- Public (unauthenticated) patient form ----------------
@router.post("/forms/bulk-delete")
async def bulk_delete_forms(data: IdList, user: dict = Depends(require_roles(*FORMS_ROLES))):
    docs = await db.forms.find({"id": {"$in": data.ids}}, {"_id": 0, "id": 1, "created_by": 1}).to_list(1000)
    deletable = [d["id"] for d in docs
                 if user["role"] == "admin" or d.get("created_by") in (user["name"], None)]
    if deletable:
        await db.forms.delete_many({"id": {"$in": deletable}})
    await log_audit("delete", "form", actor=user, detail=f"bulk ({len(deletable)})")
    return {"deleted": len(deletable), "skipped": len(data.ids) - len(deletable)}


@router.delete("/forms/{fid}")
async def delete_form(fid: str, user: dict = Depends(require_roles(*FORMS_ROLES))):
    existing = await db.forms.find_one({"id": fid}, {"_id": 0})
    if not existing:
        raise HTTPException(status_code=404, detail="Form not found")
    if user["role"] != "admin" and existing.get("created_by") not in (user["name"], None):
        raise HTTPException(status_code=403, detail="You can only delete forms you created")
    await db.forms.delete_one({"id": fid})
    await log_audit("delete", "form", actor=user, resource_id=fid, detail=existing.get("title", ""))
    return {"ok": True}


@router.get("/public/forms/{token}")
async def public_get_form(token: str):
    f = await db.forms.find_one({"public_token": token}, {"_id": 0, "created_by": 0})
    if not f:
        raise HTTPException(status_code=404, detail="Form not found")
    patient = None
    if f.get("patient_id"):
        p = await db.patients.find_one({"id": f["patient_id"]}, {"_id": 0, "first_name": 1})
        patient = p["first_name"] if p else None
    return {"id": f["id"], "title": f["title"], "form_type": f["form_type"],
            "template": f.get("template", []), "status": f["status"],
            "has_template": bool(f.get("template")), "has_attachment": bool(f.get("attachment")),
            "patient_first_name": patient, "clinic": "Veterans of Puerto Plata"}


@router.post("/public/forms/{token}/submit")
async def public_submit_form(token: str, data: FormSubmission):
    import json as _json
    if len(_json.dumps(data.responses)) > 1_200_000:
        raise HTTPException(status_code=400, detail="Submission too large")
    f = await db.forms.find_one({"public_token": token})
    if not f:
        raise HTTPException(status_code=404, detail="Form not found")
    if f.get("status") == "received":
        raise HTTPException(status_code=400, detail="This form has already been submitted")
    missing = [fld["name"] for fld in f.get("template", [])
               if fld.get("required") and not data.responses.get(fld["name"])]
    if missing:
        raise HTTPException(status_code=400, detail=f"Missing required fields: {', '.join(missing)}")
    await db.forms.update_one({"public_token": token},
        {"$set": {"responses": data.responses, "status": "received", "submitted_at": now_iso()}})
    return {"ok": True}


@router.post("/public/forms/{token}/upload")
async def public_upload_back(token: str, file: UploadFile = File(...)):
    f = await db.forms.find_one({"public_token": token})
    if not f:
        raise HTTPException(status_code=404, detail="Form not found")
    if f.get("status") == "received":
        raise HTTPException(status_code=400, detail="This form has already been submitted")
    ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else "bin"
    allowed = {"pdf", "png", "jpg", "jpeg", "webp", "doc", "docx", "txt", "xls", "xlsx"}
    if ext not in allowed:
        raise HTTPException(status_code=400, detail="Unsupported file type")
    ct_map = {"xls": "application/vnd.ms-excel",
              "xlsx": "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"}
    safe_ct = EXT_CONTENT_TYPES.get(ext) or ct_map.get(ext, "application/octet-stream")
    payload = await file.read()
    if len(payload) > 15 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 15MB)")
    path = f"{APP_NAME}/forms/patient-uploads/{token}/{uuid.uuid4()}.{ext}"
    try:
        result = put_object(path, payload, safe_ct)
    except Exception as e:
        logger.error(f"patient upload failed: {e}")
        raise HTTPException(status_code=502, detail="File storage failed")
    await db.forms.update_one({"public_token": token}, {"$set": {
        "attachment": {"storage_path": result["path"], "filename": file.filename,
                       "content_type": safe_ct, "size": result.get("size")},
        "status": "received", "submitted_at": now_iso()}})
    return {"ok": True}
